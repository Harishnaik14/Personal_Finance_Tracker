from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.db.models import Sum
from .models import Transaction, Category, Goal, Challenge, Badge
from .forms import TransactionForm, GoalForm, ExportForm
from django.contrib import messages
from django.utils import timezone
from decimal import Decimal
from datetime import timedelta
from .utils import convert_amount, get_spending_insights, get_ai_suggestions, init_guest_session, get_guest_data, add_guest_transaction

def home(request):
    return redirect('dashboard')


def dashboard(request):
    if not request.user.is_authenticated:
        # --- GUEST SESSION MODE ---
        init_guest_session(request)
        guest_data = get_guest_data(request)
        
        user_currency = request.session.get('guest_currency', 'USD')
        currency_symbol = '$'
        
        # Dynamic Data from Session
        balance = guest_data['balance']
        monthly_income = guest_data['income']
        monthly_expense = guest_data['expense']
        recent_transactions = guest_data['recent_transactions']
        
        # Demo Charts (Static for now to ensure visual consistency)
        category_labels = ['Food', 'Travel', 'Shopping', 'Bills', 'Entertainment']
        category_data = [450, 300, 250, 600, 150]
        
        months_labels = ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec']
        # Nice curve
        monthly_data = [1200, 1400, 1100, 1600, 1800, 1500, 1900, 2100, 1850.00, 0, 0, 0] 

        context = {
            'income': monthly_income,
            'expense': monthly_expense,
            'balance': balance,
            'recent_transactions': recent_transactions,
            'category_labels': category_labels,
            'category_data': category_data,
            'months_labels': months_labels,
            'monthly_data': monthly_data,
            'current_year': timezone.now().year,
            'currency_symbol': currency_symbol,
            'is_demo': True, # Flag for template
            'user_currency': user_currency, # Pass explicit currency for templates
        }
        
        # Demo Insights
        context['spending_insights'] = [
            {'category': 'Food', 'status': 'increase', 'text': 'Spending is 15% higher than last month.'},
            {'category': 'Travel', 'status': 'decrease', 'text': 'Spending is 10% lower than last month.'},
        ]
        context['total_insight'] = {'status': 'increase', 'text': 'Total spending is on track.'}
        
        # Demo Goals
        context['top_goals'] = [
            {'name': 'New Laptop', 'percent': 75, 'status_color': 'warning'},
            {'name': 'Vacation', 'percent': 40, 'status_color': 'danger'},
            {'name': 'Emergency Fund', 'percent': 90, 'status_color': 'success'},
        ]
        
        # Demo Challenge
        context['active_challenge'] = {
            'obj': {'get_challenge_type_display': 'No Spend Month'}, # Mock object
            'percent': 65,
            'days_left': 10
        }

        return render(request, 'finance/dashboard.html', context)

    # --- AUTHENTICATED USER LOGIC ---
    # Summary Logic
    transactions = Transaction.objects.filter(user=request.user).order_by('-date', '-created_at')
    
    now = timezone.now()
    current_month_transactions = transactions.filter(date__year=now.year, date__month=now.month)
    
    # Current Month Stats
    monthly_expense = current_month_transactions.filter(category__type='expense').aggregate(Sum('amount'))['amount__sum'] or 0
    monthly_income = current_month_transactions.filter(category__type='income').aggregate(Sum('amount'))['amount__sum'] or 0
    
    # Total Balance Logic (All time) - includes opening balance
    total_income = Transaction.objects.filter(user=request.user, category__type='income').aggregate(Sum('amount'))['amount__sum'] or 0
    total_expense = Transaction.objects.filter(user=request.user, category__type='expense').aggregate(Sum('amount'))['amount__sum'] or 0
    balance = float(request.user.opening_balance) + float(total_income) - float(total_expense)

    # Alert user to set opening balance if not set (only show once per session)
    if request.user.opening_balance == 0 and 'opening_balance_alert_shown' not in request.session:
        from django.contrib import messages
        messages.info(request, 'Welcome! Set your opening balance in Settings to track your finances accurately.', extra_tags='opening_balance_alert')
        request.session['opening_balance_alert_shown'] = True

    recent_transactions = transactions[:10]

    # --- CURRENCY CONVERSION ---
    user_currency = request.user.currency
    
    # Convert Scalars
    display_expense = convert_amount(monthly_expense, user_currency)
    display_income = convert_amount(monthly_income, user_currency)
    display_balance = convert_amount(balance, user_currency)

    # Chart Data 1: Expenses by Category (Doughnut)
    expense_by_category = current_month_transactions.filter(category__type='expense').values('category__name').annotate(total=Sum('amount')).order_by('-total')
    category_labels = [item['category__name'] for item in expense_by_category]
    category_data = [convert_amount(item['total'], user_currency) for item in expense_by_category]

    # Chart Data 2: Expenses by Month (Bar Chart) for Current Year
    from django.db.models.functions import ExtractMonth
    monthly_stats = transactions.filter(
        date__year=now.year, 
        category__type='expense'
    ).annotate(month=ExtractMonth('date')).values('month').annotate(total=Sum('amount')).order_by('month')

    monthly_data = [0] * 12
    for item in monthly_stats:
        monthly_data[item['month'] - 1] = convert_amount(float(item['total']), user_currency)
    
    months_labels = ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec']


    context = {
        'income': display_income,
        'expense': display_expense,
        'balance': display_balance,
        'recent_transactions': recent_transactions,
        'category_labels': category_labels,
        'category_data': category_data,
        'months_labels': months_labels,
        'monthly_data': monthly_data,
        'current_year': now.year,
        'is_demo': False,
        'user_currency': user_currency
    }

    # Spending Insights
    insights_data = get_spending_insights(request.user)
    context.update({
        'spending_insights': insights_data['category_insights'][:5], 
        'total_insight': insights_data['total_insight'],
        'has_prev_data': insights_data['has_prev_data']
    })

    # Top Goals Summary
    goals = Goal.objects.filter(user=request.user)[:3]
    top_goals = []
    for goal in goals:
        saved = goal.saved_amount
        percent = min(100, round((float(saved) / float(goal.target_amount)) * 100, 1)) if goal.target_amount > 0 else 0
        top_goals.append({
            'name': goal.name,
            'percent': percent,
            'status_color': 'danger' if percent < 40 else ('warning' if percent < 80 else 'success')
        })
    context['top_goals'] = top_goals

    # Active Challenge Summary
    active_challenge = Challenge.objects.filter(user=request.user, is_active=True).first()
    if active_challenge:
        total_days = (active_challenge.end_date - active_challenge.start_date).days + 1
        elapsed_days = (now.date() - active_challenge.start_date).days
        context['active_challenge'] = {
            'obj': active_challenge,
            'percent': min(100, round((elapsed_days / total_days) * 100)),
            'days_left': max(0, (active_challenge.end_date - now.date()).days)
        }

    return render(request, 'finance/dashboard.html', context)


@login_required
def simulate_payment(request):
    import random
    if request.method == 'POST':
        # Simulate fetching data from API
        vendors = ['Uber', 'Swiggy', 'Amazon', 'Starbucks', 'Netflix', 'Freelance Client', 'Monthly Salary', 'Cashback']
        
        count = 0
        for _ in range(4): # Add 4 random transactions
            vendor = random.choice(vendors)
            
            if vendor in ['Freelance Client', 'Monthly Salary', 'Cashback']:
                # Income
                cat_name = 'Salary' if vendor == 'Monthly Salary' else 'Other'
                if vendor == 'Freelance Client': cat_name = 'Freelance'
                
                category = Category.objects.filter(name=cat_name, type='income').first()
                if not category:
                    category = Category.objects.filter(type='income').first()
                
                amount = random.randint(100, 1000) * 10 if vendor != 'Cashback' else random.randint(1, 50)
                description = f"Received {vendor} payment"
            else:
                # Expense
                cat_name = 'Other'
                if vendor == 'Uber': cat_name = 'Travel'
                elif vendor == 'Swiggy': cat_name = 'Food'
                elif vendor == 'Amazon': cat_name = 'Shopping'
                elif vendor == 'Starbucks': cat_name = 'Food'
                elif vendor == 'Netflix': cat_name = 'Entertainment'
                
                category = Category.objects.filter(name=cat_name, type='expense').first()
                if not category:
                    category = Category.objects.filter(type='expense').first()
                    
                amount = random.randint(5, 50) * 10
                description = f"Payment to {vendor} (via UPI)"
            
            Transaction.objects.create(
                user=request.user,
                category=category,
                amount=amount,
                date=timezone.now(),
                description=description
            )
            count += 1
            
        messages.success(request, f"Successfully imported {count} transactions (Income & Expenses)!")
        return redirect('dashboard')
        
    return render(request, 'finance/simulate_payment.html')

@login_required # Keep this for now or handle guest?
def transaction_list(request):
    # For now, just keep it login protected to avoid complexity with 'years', 'months' logic which is heavy on DB queries.
    # The user wanted "Add Transaction" to work.
    # If they click "View All" on dashboard, it takes them here.
    # I should Redirect guest to dashboard with a message "Full history available after login"?
    # OR implement a simple guest list.
    if not request.user.is_authenticated:
         # Simplified Guest List
         init_guest_session(request)
         guest_data = get_guest_data(request)
         
         # Mocking year context for template
         return render(request, 'finance/transaction_list.html', {
            'transactions': guest_data['all_transactions'], # List of dicts
            'current_year': timezone.now().year,
            'current_month': 'all',
            'years_list': [2024, 2023],
            'months': [{'num': i, 'name': n, 'disabled': False} for i, n in enumerate(['Jan', 'Feb', 'Mar'], 1)], # Dummy
            'current_sort': 'category',
            'is_guest': True,
            'user_currency': 'USD',
            'currency_symbol': '$'
         })

    now = timezone.now()
    try:
        current_year = int(request.GET.get('year', now.year))
    except ValueError:
        current_year = now.year

    # Filter transactions by year (optional - if user wants only current year by default)
    transactions = Transaction.objects.filter(user=request.user, date__year=current_year)

    # Filter by Month
    current_month = request.GET.get('month', 'all')
    if current_month != 'all':
        try:
            month_int = int(current_month)
            transactions = transactions.filter(date__month=month_int)
        except ValueError:
            current_month = 'all' # Reset if invalid

    # Sorting
    sort_by = request.GET.get('sort', 'category') # Default to Category wise as per user preference
    if sort_by == 'category':
        transactions = transactions.order_by('category__name', '-date') # Group by category, then new first
    else:
        transactions = transactions.order_by('-date', '-created_at')

    # Year Range: Current Year down to 2025 PLUS any transaction years
    fixed_range = list(range(now.year, 2024, -1))
    # Available years from transactions
    available_years = Transaction.objects.filter(user=request.user).dates('date', 'year')
    trans_years = [d.year for d in available_years]
    
    # Merge and sort
    years_list = sorted(list(set(fixed_range + trans_years)), reverse=True)

    months_list = [
        {'num': 1, 'name': 'Jan'}, {'num': 2, 'name': 'Feb'}, {'num': 3, 'name': 'Mar'},
        {'num': 4, 'name': 'Apr'}, {'num': 5, 'name': 'May'}, {'num': 6, 'name': 'Jun'},
        {'num': 7, 'name': 'Jul'}, {'num': 8, 'name': 'Aug'}, {'num': 9, 'name': 'Sep'},
        {'num': 10, 'name': 'Oct'}, {'num': 11, 'name': 'Nov'}, {'num': 12, 'name': 'Dec'}
    ]
    
    # Logic to disable months prior to registration, UNLESS they have data
    date_joined = request.user.date_joined
    join_year = date_joined.year
    join_month = date_joined.month
    
    # Get months with actual transactions for this year
    # Aggregating to find which months have data
    from django.db.models.functions import ExtractMonth
    active_months_qs = Transaction.objects.filter(user=request.user, date__year=current_year).annotate(m=ExtractMonth('date')).values_list('m', flat=True).distinct()
    active_months = set(active_months_qs)

    months = []
    for m in months_list:
        is_disabled = False
        # Rule: Disable if year < join_year OR (year == join_year and month < join_month)
        # BUT: Enable if active_months has data (user manually added backdated transaction)
        if current_year < join_year:
            is_disabled = True
        elif current_year == join_year and m['num'] < join_month:
            is_disabled = True
            
        # Exception: Allow if data exists
        if m['num'] in active_months:
            is_disabled = False
            
        m['disabled'] = is_disabled
        months.append(m)

    return render(request, 'finance/transaction_list.html', {
        'transactions': transactions,
        'current_year': current_year,
        'current_month': current_month if current_month == 'all' else int(current_month),
        'years_list': years_list,
        'months': months,
        'current_sort': sort_by,
        'user_currency': request.user.currency
    })

def transaction_create(request):
    if request.method == 'POST':
        if request.user.is_authenticated:
            form = TransactionForm(request.user, request.POST)
            if form.is_valid():
                transaction = form.save(commit=False)
                transaction.user = request.user
                
                # --- CHALLENGE VIOLATION CHECK ---
                if transaction.category.type == 'expense':
                    active_challenge = Challenge.objects.filter(user=request.user, is_active=True).first()
                    if active_challenge:
                        # Check if transaction date is within challenge period
                        # AND ensure transaction was created AFTER the challenge started
                        if active_challenge.start_date <= transaction.date <= active_challenge.end_date and transaction.created_at >= active_challenge.created_at:
                            active_challenge.is_active = False
                            active_challenge.is_successful = False
                            active_challenge.save()
                            messages.warning(request, f'Challenge Failed! Spending recorded on {transaction.date} violated your {active_challenge.challenge_type} challenge.')

                transaction.save()
                messages.success(request, 'Transaction added successfully!')
                # Use reverse to look up the correct URL path dynamically
                from django.urls import reverse
                base_url = reverse('transaction_list')
                return redirect(f"{base_url}?year={transaction.date.year}&month={transaction.date.month}")
        else:
            # Guest Logic
            form = TransactionForm(user=None, data=request.POST)
            if form.is_valid():
                add_guest_transaction(request, form.cleaned_data)
                messages.success(request, 'Transaction added to guest session!')
                return redirect('dashboard')

    else:
        # GET
        if request.user.is_authenticated:
            form = TransactionForm(request.user)
        else:
            form = TransactionForm(user=None)
            
    return render(request, 'finance/transaction_form.html', {'form': form, 'title': 'Add Transaction'})

@login_required
def transaction_update(request, pk):
    transaction = get_object_or_404(Transaction, pk=pk, user=request.user)
    if request.method == 'POST':
        form = TransactionForm(request.user, request.POST, instance=transaction)
        if form.is_valid():
            form.save()
            messages.success(request, 'Transaction updated!')
            return redirect('transaction_list')
    else:
        form = TransactionForm(request.user, instance=transaction)
    return render(request, 'finance/transaction_form.html', {'form': form, 'title': 'Edit Transaction'})

@login_required
def overview(request):
    from django.db.models.functions import ExtractMonth
    
    # Year Handling
    now = timezone.now()
    try:
        current_year = int(request.GET.get('year', now.year))
    except ValueError:
        current_year = now.year

    # Filter transactions by year
    transactions = Transaction.objects.filter(user=request.user, date__year=current_year)

    # 1. Monthly Data
    monthly_stats = transactions.values('category__type', 'date__month').annotate(total=Sum('amount')).order_by('date__month')
    
    monthly_income_data = [0] * 12
    monthly_expense_data = [0] * 12
    user_currency = request.user.currency

    for item in monthly_stats:
        month_idx = item['date__month'] - 1
        amount = float(item['total'])
        converted_amount = convert_amount(amount, user_currency)
        
        if item['category__type'] == 'income':
            monthly_income_data[month_idx] = converted_amount
        else:
            monthly_expense_data[month_idx] = converted_amount
    
    months_labels = ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec']

    # Chart Data 2: Expenses by Category (Pie Chart) for Selected Year
    expense_category_qs = transactions.filter(category__type='expense').values('category__name').annotate(total=Sum('amount')).order_by('-total')
    pie_labels = [item['category__name'] for item in expense_category_qs]
    pie_data = [convert_amount(item['total'], user_currency) for item in expense_category_qs]

    # Get available years for dropdown
    available_years = Transaction.objects.filter(user=request.user).dates('date', 'year')
    years_list = sorted(list(set([d.year for d in available_years] + list(range(now.year, 2024, -1)))), reverse=True)

    context = {
        'months_labels': months_labels,
        'monthly_income_data': monthly_income_data,
        'monthly_expense_data': monthly_expense_data,
        'pie_labels': pie_labels,
        'pie_data': pie_data,
        'current_year': current_year,
        'years_list': years_list,
        'transactions': transactions.order_by('category__name', '-date')
    }
    return render(request, 'finance/overview.html', context)

@login_required
def transaction_delete(request, pk):
    transaction = get_object_or_404(Transaction, pk=pk, user=request.user)
    if request.method == 'POST':
        transaction.delete()
        messages.success(request, 'Transaction deleted!')
        return redirect('transaction_list')
    return render(request, 'finance/transaction_confirm_delete.html', {'transaction': transaction})

@login_required
def export_data(request):
    import csv
    from django.http import HttpResponse
    from django.utils.text import slugify
    from io import BytesIO
    
    preview_data = None
    transactions_list = None
    form = ExportForm(request.GET or None)

    if form.is_valid():
        file_name = form.cleaned_data['file_name']
        start_date = form.cleaned_data['start_date']
        end_date = form.cleaned_data['end_date']
        file_format = form.cleaned_data['format']
        include_transactions = form.cleaned_data.get('include_transactions', True)

        # Filter transactions
        transactions = Transaction.objects.filter(
            user=request.user,
            date__range=[start_date, end_date]
        ).order_by('date')

        user_currency = request.user.currency

        # Check if user clicked Download
        if 'download' in request.GET:
            # Smart filename handling - check if extension already exists
            file_name_lower = file_name.lower()
            
            if file_format == 'csv':
                # Only add .csv if not already present
                if not file_name_lower.endswith('.csv'):
                    filename = f"{slugify(file_name)}.csv"
                else:
                    filename = slugify(file_name)
                
                response = HttpResponse(content_type='text/csv')
                response['Content-Disposition'] = f'attachment; filename="{filename}"'

                writer = csv.writer(response)
                writer.writerow(['Date', 'Category', 'Type', 'Description', 'Amount', 'Currency'])

                for trans in transactions:
                    writer.writerow([
                        trans.date,
                        trans.category.name,
                        trans.category.type,
                        trans.description,
                        trans.amount,
                        user_currency
                    ])

                return response
            
            elif file_format == 'pdf':
                from reportlab.lib.pagesizes import letter, A4
                from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
                from reportlab.lib import colors
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.units import inch
                from reportlab.lib.enums import TA_CENTER
                
                buffer = BytesIO()
                doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=0.5*inch)
                elements = []
                styles = getSampleStyleSheet()
                
                # Title
                title_style = ParagraphStyle(
                    'CustomTitle',
                    parent=styles['Heading1'],
                    fontSize=24,
                    textColor=colors.HexColor('#667eea'),
                    spaceAfter=20,
                    alignment=TA_CENTER
                )
                title = Paragraph(f"<b>{file_name}</b>", title_style)
                elements.append(title)
                
                # Period
                period_style = ParagraphStyle('Period', parent=styles['Normal'], alignment=TA_CENTER, fontSize=11)
                period = Paragraph(f"<b>Period:</b> {start_date.strftime('%B %d, %Y')} - {end_date.strftime('%B %d, %Y')}", period_style)
                elements.append(period)
                elements.append(Spacer(1, 0.3*inch))
                
                # Calculate totals
                income_total = transactions.filter(category__type='income').aggregate(Sum('amount'))['amount__sum'] or 0
                expense_total = transactions.filter(category__type='expense').aggregate(Sum('amount'))['amount__sum'] or 0
                income_converted = convert_amount(float(income_total), user_currency)
                expense_converted = convert_amount(float(expense_total), user_currency)
                balance = income_converted - expense_converted
                
                # Create colorful summary boxes
                summary_data = [
                    ['Total Income', 'Total Expenses'],
                    [f'{request.user.currency_symbol}{income_converted:.2f}',
                     f'{request.user.currency_symbol}{expense_converted:.2f}']
                ]
                
                summary_table = Table(summary_data, colWidths=[3.75*inch, 3.75*inch])
                summary_table.setStyle(TableStyle([
                    # Header row with gradient colors
                    ('BACKGROUND', (0, 0), (0, 0), colors.HexColor('#11998e')),  # Green for Income
                    ('BACKGROUND', (1, 0), (1, 0), colors.HexColor('#ee0979')),  # Red for Expenses
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 11),
                    ('FONTSIZE', (0, 1), (-1, 1), 16),
                    ('FONTNAME', (0, 1), (-1, 1), 'Helvetica-Bold'),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('TOPPADDING', (0, 1), (-1, 1), 15),
                    ('BOTTOMPADDING', (0, 1), (-1, 1), 15),
                    ('BOX', (0, 0), (-1, -1), 2, colors.white),
                    ('LINEABOVE', (0, 1), (-1, 1), 2, colors.lightgrey),
                    ('BACKGROUND', (0, 1), (-1, 1), colors.HexColor('#f8f9fa')),
                ]))
                elements.append(summary_table)
                elements.append(Spacer(1, 0.4*inch))
                
                # Month-over-Month Comparison (if available)
                from dateutil.relativedelta import relativedelta
                prev_month_start = start_date - relativedelta(months=1)
                prev_month_end = end_date - relativedelta(months=1)
                prev_transactions = Transaction.objects.filter(
                    user=request.user,
                    date__range=[prev_month_start, prev_month_end]
                )
                prev_expense_total = prev_transactions.filter(category__type='expense').aggregate(Sum('amount'))['amount__sum'] or 0
                prev_expense_converted = convert_amount(float(prev_expense_total), user_currency)
                
                if prev_expense_converted >= 0:
                    expense_change_pct = ((expense_converted - prev_expense_converted) / prev_expense_converted * 100) if prev_expense_converted > 0 else 0
                    
                    comparison_header = Paragraph("<b>Month-over-Month Comparison</b>", styles['Heading3'])
                    elements.append(comparison_header)
                    elements.append(Spacer(1, 0.1*inch))
                    
                    comparison_data = [
                        ['Previous Month Expenses', 'Change'],
                        [f'{request.user.currency_symbol}{prev_expense_converted:.2f}',
                         f"{'↑' if expense_change_pct > 0 else '↓'} {abs(expense_change_pct):.1f}%"]
                    ]
                    
                    comparison_table = Table(comparison_data, colWidths=[3.75*inch, 3.75*inch])
                    comparison_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f093fb')),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 11),
                        ('FONTSIZE', (0, 1), (-1, 1), 14),
                        ('FONTNAME', (0, 1), (-1, 1), 'Helvetica-Bold'),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                        ('TOPPADDING', (0, 1), (-1, 1), 12),
                        ('BOTTOMPADDING', (0, 1), (-1, 1), 12),
                        ('BOX', (0, 0), (-1, -1), 2, colors.white),
                        ('BACKGROUND', (0, 1), (-1, 1), colors.HexColor('#fce4ec')),
                    ]))
                    elements.append(comparison_table)
                    elements.append(Spacer(1, 0.3*inch))
                
                # Category Breakdown Charts
                from reportlab.graphics.shapes import Drawing
                from reportlab.graphics.charts.piecharts import Pie
                from reportlab.graphics.charts.legends import Legend
                from reportlab.lib import colors as rl_colors
                
                expense_by_category = transactions.filter(category__type='expense').values('category__name').annotate(total=Sum('amount')).order_by('-total')[:5]
                income_by_category = transactions.filter(category__type='income').values('category__name').annotate(total=Sum('amount')).order_by('-total')[:5]
                
                if expense_by_category or income_by_category:
                    chart_header = Paragraph("<b>Breakdown by Category</b>", styles['Heading3'])
                    elements.append(chart_header)
                    elements.append(Spacer(1, 0.2*inch))
                    
                    # Expense Chart
                    expense_drawing = None
                    if expense_by_category:
                        expense_drawing = Drawing(250, 250)
                        expense_pie = Pie()
                        expense_pie.x = 75
                        expense_pie.y = 90
                        expense_pie.width = 100
                        expense_pie.height = 100
                        
                        expense_data = [float(convert_amount(float(item['total']), user_currency)) for item in expense_by_category]
                        expense_labels = [item['category__name'][:12] for item in expense_by_category]
                        
                        expense_pie.data = expense_data
                        expense_pie.labels = []  # Remove labels from pie, use legend instead
                        expense_pie.slices.strokeColor = rl_colors.white
                        expense_pie.slices.strokeWidth = 2
                        
                        # Red shades for expenses
                        expense_colors = [
                            rl_colors.HexColor('#ef4444'),
                            rl_colors.HexColor('#f97316'),
                            rl_colors.HexColor('#fb923c'),
                            rl_colors.HexColor('#fbbf24'),
                            rl_colors.HexColor('#fcd34d')
                        ]
                        for i, color in enumerate(expense_colors[:len(expense_data)]):
                            expense_pie.slices[i].fillColor = color
                        
                        # Add legend
                        legend = Legend()
                        legend.x = 10
                        legend.y = 20
                        legend.dx = 8
                        legend.dy = 8
                        legend.fontSize = 8
                        legend.columnMaximum = 5
                        legend.colorNamePairs = [(expense_colors[i], expense_labels[i]) for i in range(len(expense_labels))]
                        
                        # Add title
                        from reportlab.graphics.shapes import String
                        title = String(125, 210, 'TOP EXPENSES', fontSize=10, fontName='Helvetica-Bold', textAnchor='middle')
                        
                        expense_drawing.add(expense_pie)
                        expense_drawing.add(legend)
                        expense_drawing.add(title)
                        
                    # Income Chart  
                    income_drawing = None
                    if income_by_category:
                        income_drawing = Drawing(250, 250)
                        income_pie = Pie()
                        income_pie.x = 75
                        income_pie.y = 90
                        income_pie.width = 100
                        income_pie.height = 100
                        
                        income_data = [float(convert_amount(float(item['total']), user_currency)) for item in income_by_category]
                        income_labels = [item['category__name'][:12] for item in income_by_category]
                        
                        income_pie.data = income_data
                        income_pie.labels = []  # Remove labels from pie, use legend instead
                        income_pie.slices.strokeColor = rl_colors.white
                        income_pie.slices.strokeWidth = 2
                        
                        # Green shades for income
                        income_colors = [
                            rl_colors.HexColor('#10b981'),
                            rl_colors.HexColor('#34d399'),
                            rl_colors.HexColor('#6ee7b7'),
                            rl_colors.HexColor('#a7f3d0'),
                            rl_colors.HexColor('#d1fae5')
                        ]
                        for i, color in enumerate(income_colors[:len(income_data)]):
                            income_pie.slices[i].fillColor = color
                        
                        # Add legend
                        legend = Legend()
                        legend.x = 10
                        legend.y = 20
                        legend.dx = 8
                        legend.dy = 8
                        legend.fontSize = 8
                        legend.columnMaximum = 5
                        legend.colorNamePairs = [(income_colors[i], income_labels[i]) for i in range(len(income_labels))]
                        
                        # Add title
                        from reportlab.graphics.shapes import String
                        title = String(125, 210, 'TOP INCOME', fontSize=10, fontName='Helvetica-Bold', textAnchor='middle')
                        
                        income_drawing.add(income_pie)
                        income_drawing.add(legend)
                        income_drawing.add(title)
                    
                    # Add charts in a table layout
                    if expense_by_category and income_by_category:
                        chart_table = Table([[expense_drawing, income_drawing]], colWidths=[3.75*inch, 3.75*inch])
                    elif expense_by_category:
                        chart_table = Table([[expense_drawing]], colWidths=[7.5*inch])
                    else:
                        chart_table = Table([[income_drawing]], colWidths=[7.5*inch])
                    
                    chart_table.setStyle(TableStyle([
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ]))
                    elements.append(chart_table)
                    elements.append(Spacer(1, 0.3*inch))
                
                # Transactions Table (only if include_transactions is True)
                if include_transactions:
                    # Section header
                    trans_header = Paragraph("<b>Transaction Details</b>", styles['Heading2'])
                    elements.append(trans_header)
                    elements.append(Spacer(1, 0.2*inch))
                    
                    data = [['Date', 'Category', 'Type', 'Description', 'Amount']]
                    for trans in transactions:
                        amount_converted = convert_amount(float(trans.amount), user_currency)
                        data.append([
                            trans.date.strftime('%Y-%m-%d'),
                            trans.category.name[:15],
                            trans.category.type.capitalize(),
                            trans.description[:30],
                            f'{request.user.currency_symbol}{amount_converted:.2f}'
                        ])
                    
                    table = Table(data, colWidths=[1.2*inch, 1.3*inch, 0.9*inch, 2*inch, 1.2*inch])
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#667eea')),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 10),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8f9fa')])
                    ]))
                    elements.append(table)
                
                doc.build(elements)
                buffer.seek(0)
                
                # Only add .pdf if not already present
                if not file_name_lower.endswith('.pdf'):
                    filename = f"{slugify(file_name)}.pdf"
                else:
                    filename = slugify(file_name)
                
                response = HttpResponse(buffer, content_type='application/pdf')
                response['Content-Disposition'] = f'attachment; filename="{filename}"'
                return response
            
            elif file_format == 'word':
                from docx import Document
                from docx.shared import Inches, Pt, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                
                doc = Document()
                
                # Title
                title = doc.add_heading(file_name, 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Period
                period_p = doc.add_paragraph(f"Period: {start_date.strftime('%B %d, %Y')} - {end_date.strftime('%B %d, %Y')}")
                period_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                doc.add_paragraph()
                
                # Summary
                doc.add_heading('Financial Summary', level=2)
                income_total = transactions.filter(category__type='income').aggregate(Sum('amount'))['amount__sum'] or 0
                expense_total = transactions.filter(category__type='expense').aggregate(Sum('amount'))['amount__sum'] or 0
                income_converted = convert_amount(float(income_total), user_currency)
                expense_converted = convert_amount(float(expense_total), user_currency)
                
                summary_table = doc.add_table(rows=4, cols=2)
                summary_table.style = 'Light Grid Accent 1'
                
                summary_table.cell(0, 0).text = 'Metric'
                summary_table.cell(0, 1).text = 'Value'
                summary_table.cell(1, 0).text = 'Total Income'
                summary_table.cell(1, 1).text = f'{request.user.currency_symbol}{income_converted:.2f}'
                summary_table.cell(2, 0).text = 'Total Expenses'
                summary_table.cell(2, 1).text = f'{request.user.currency_symbol}{expense_converted:.2f}'
                summary_table.cell(3, 0).text = 'Net Balance'
                summary_table.cell(3, 1).text = f'{request.user.currency_symbol}{(income_converted - expense_converted):.2f}'
                
                doc.add_paragraph()
                
                # Transactions (only if include_transactions is True)
                if include_transactions:
                    doc.add_heading('Transaction Details', level=2)
                    trans_table = doc.add_table(rows=1, cols=5)
                    trans_table.style = 'Light Grid Accent 1'
                    
                    hdr_cells = trans_table.rows[0].cells
                    hdr_cells[0].text = 'Date'
                    hdr_cells[1].text = 'Category'
                    hdr_cells[2].text = 'Type'
                    hdr_cells[3].text = 'Description'
                    hdr_cells[4].text = 'Amount'
                    
                    for trans in transactions:
                        row_cells = trans_table.add_row().cells
                        amount_converted = convert_amount(float(trans.amount), user_currency)
                        row_cells[0].text = trans.date.strftime('%Y-%m-%d')
                        row_cells[1].text = trans.category.name
                        row_cells[2].text = trans.category.type.capitalize()
                        row_cells[3].text = trans.description
                        row_cells[4].text = f'{request.user.currency_symbol}{amount_converted:.2f}'
                
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                
                # Only add .docx if not already present
                if not file_name_lower.endswith('.docx') and not file_name_lower.endswith('.doc'):
                    filename = f"{slugify(file_name)}.docx"
                else:
                    filename = slugify(file_name)
                
                response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Disposition'] = f'attachment; filename="{filename}"'
                return response

        # Otherwise, provide Preview data with full transaction list
        income_total = transactions.filter(category__type='income').aggregate(Sum('amount'))['amount__sum'] or 0
        expense_total = transactions.filter(category__type='expense').aggregate(Sum('amount'))['amount__sum'] or 0
        
        # Calculate previous month for comparison
        from dateutil.relativedelta import relativedelta
        prev_month_start = start_date - relativedelta(months=1)
        prev_month_end = end_date - relativedelta(months=1)
        
        prev_transactions = Transaction.objects.filter(
            user=request.user,
            date__range=[prev_month_start, prev_month_end]
        )
        prev_expense_total = prev_transactions.filter(category__type='expense').aggregate(Sum('amount'))['amount__sum'] or 0
        prev_expense_converted = convert_amount(float(prev_expense_total), user_currency)
        
        # Category breakdown for chart
        expense_by_category = transactions.filter(category__type='expense').values('category__name').annotate(total=Sum('amount')).order_by('-total')[:5]
        income_by_category = transactions.filter(category__type='income').values('category__name').annotate(total=Sum('amount')).order_by('-total')[:5]
        
        expense_categories = [item['category__name'] for item in expense_by_category]
        expense_amounts = [float(convert_amount(float(item['total']), user_currency)) for item in expense_by_category]
        
        income_categories = [item['category__name'] for item in income_by_category]
        income_amounts = [float(convert_amount(float(item['total']), user_currency)) for item in income_by_category]
        
        # Calculate balance
        income_converted = convert_amount(float(income_total), user_currency)
        expense_converted = convert_amount(float(expense_total), user_currency)
        balance = income_converted - expense_converted
        
        # Calculate month-over-month change
        if prev_expense_converted > 0:
            expense_change_pct = ((expense_converted - prev_expense_converted) / prev_expense_converted) * 100
        else:
            expense_change_pct = 0 if expense_converted == 0 else 100
        
        preview_data = {
            'income': income_converted,
            'expense': expense_converted,
            'balance': balance,
            'count': transactions.count(),
            'start': start_date,
            'end': end_date,
            'format': file_format,
            'include_transactions': include_transactions,
            'expense_categories': expense_categories,
            'expense_amounts': expense_amounts,
            'income_categories': income_categories,
            'income_amounts': income_amounts,
            'prev_expense': prev_expense_converted,
            'expense_change_pct': expense_change_pct,
            'expense_increased': expense_change_pct > 0
        }
        
        # Full transaction list for review (only if requested)
        if include_transactions:
            transactions_list = []
            for trans in transactions:
                transactions_list.append({
                    'date': trans.date,
                    'category': trans.category.name,
                    'type': trans.category.type,
                    'description': trans.description,
                    'amount': convert_amount(float(trans.amount), user_currency)
                })

    return render(request, 'finance/export_data.html', {
        'form': form,
        'preview': preview_data,
        'transactions': transactions_list
    })

@login_required
def calendar_view(request):
    import calendar as py_calendar
    
    # Use Local Time for calculations
    now = timezone.localtime(timezone.now())
    
    try:
        year = int(request.GET.get('year', now.year))
        month = int(request.GET.get('month', now.month))
    except (ValueError, TypeError):
        year, month = now.year, now.month

    # Handle month navigation overflow/underflow
    if month < 1: 
        month = 12
        year -= 1
    elif month > 12: 
        month = 1
        year += 1

    # Calendar Grid Data
    cal = py_calendar.Calendar(firstweekday=py_calendar.SUNDAY)
    month_days = cal.monthdayscalendar(year, month)
    
    # Get Transactions for this month
    transactions = Transaction.objects.filter(user=request.user, date__year=year, date__month=month)
    
    # Aggregate by day
    daily_stats = {}
    user_currency = request.user.currency
    
    for trans in transactions:
        day = trans.date.day
        if day not in daily_stats:
            daily_stats[day] = {'income': 0, 'expense': 0}
        
        amount = convert_amount(float(trans.amount), user_currency)
        if trans.category.type == 'income':
            daily_stats[day]['income'] += amount
        else:
            daily_stats[day]['expense'] += amount

    # Navigation Labels & Dates
    month_name = py_calendar.month_name[month]
    prev_month = month - 1
    prev_year = year
    if prev_month < 1: prev_month = 12; prev_year -= 1
    
    next_month = month + 1
    next_year = year
    if next_month > 12: next_month = 1; next_year += 1

    # Calculated Totals for footer
    total_income = sum(stats['income'] for stats in daily_stats.values())
    total_expense = sum(stats['expense'] for stats in daily_stats.values())

    # Build years_list to match overview condition
    available_years = Transaction.objects.filter(user=request.user).dates('date', 'year')
    years_list = sorted(list(set([d.year for d in available_years] + list(range(now.year, 2024, -1)))), reverse=True)

    months_list = [
        (1, 'January'), (2, 'February'), (3, 'March'), (4, 'April'),
        (5, 'May'), (6, 'June'), (7, 'July'), (8, 'August'),
        (9, 'September'), (10, 'October'), (11, 'November'), (12, 'December')
    ]

    context = {
        'year': year,
        'month': month,
        'month_name': month_name,
        'month_days': month_days,
        'daily_stats': daily_stats,
        'prev_month': prev_month,
        'prev_year': prev_year,
        'next_month': next_month,
        'next_year': next_year,
        'income': total_income,
        'expense': total_expense,
        'current_date': now.date(),
        'years_list': years_list,
        'months_list': months_list
    }
    return render(request, 'finance/calendar.html', context)

@login_required
def income_vs_expense(request):
    from django.db.models.functions import ExtractMonth
    
    # Year Handling
    now = timezone.now()
    try:
        current_year = int(request.GET.get('year', now.year))
    except ValueError:
        current_year = now.year

    # 1. Monthly Data
    transactions = Transaction.objects.filter(user=request.user, date__year=current_year)
    monthly_stats = transactions.values('category__type', 'date__month').annotate(total=Sum('amount')).order_by('date__month')
    
    income_data = [0] * 12
    expense_data = [0] * 12
    user_currency = request.user.currency

    for item in monthly_stats:
        month_idx = item['date__month'] - 1
        amount = convert_amount(float(item['total']), user_currency)
        if item['category__type'] == 'income':
            income_data[month_idx] = amount
        else:
            expense_data[month_idx] = amount

    months_labels = ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec']

    # Year Range: Current Year down to 2025 PLUS any transaction years
    fixed_range = list(range(now.year, 2024, -1))
    available_years = Transaction.objects.filter(user=request.user).dates('date', 'year')
    years_list = sorted(list(set(fixed_range + [d.year for d in available_years])), reverse=True)

    # AI Suggestions (Optional but nice to keep)
    ai_data = get_ai_suggestions(request.user)
    
    context = {
        'current_year': current_year,
        'years_list': years_list,
        'months_labels': months_labels,
        'income_data': income_data,
        'expense_data': expense_data,
        # AI Data
        'ai_suggestions': ai_data['suggestions'],
        'prediction_total': ai_data['prediction_total'],
        'prediction_labels': ai_data['chart_labels'],
        'prediction_data': ai_data['chart_data'],
        'historical_avg': ai_data['historical_avg']
    }
    return render(request, 'finance/income_vs_expense.html', context)
@login_required
def goal_list(request):
    if request.method == 'POST':
        form = GoalForm(request.POST)
        if form.is_valid():
            goal = form.save(commit=False)
            goal.user = request.user
            goal.save()
            messages.success(request, 'Goal created successfully!')
            return redirect('goal_list')
    else:
        form = GoalForm()

    goals = Goal.objects.filter(user=request.user)
    goals_data = []

    for goal in goals:
        saved_amount = goal.saved_amount
        percent = min(100, round((float(saved_amount) / float(goal.target_amount)) * 100, 1)) if goal.target_amount > 0 else 0
        
        # Savings Breakdown
        days_left = None
        savings_breakdown = None
        if goal.target_date:
            today = timezone.now().date()
            days_left = (goal.target_date - today).days
            if days_left > 0:
                remaining = float(goal.target_amount - saved_amount)
                if remaining > 0:
                    daily = remaining / days_left
                    monthly = remaining / (days_left / 30.44) # Average month length
                    quarterly = remaining / (days_left / 91.25) # Quarter (3 months)
                    yearly = remaining / (days_left / 365.25) if days_left > (5 * 365.25) else None
                    
                    savings_breakdown = {
                        'daily': round(daily, 2),
                        'monthly': round(monthly, 2),
                        'quarterly': round(quarterly, 2),
                        'yearly': round(yearly, 2) if yearly else None,
                        'total_years': round(days_left / 365.25, 1)
                    }

        goals_data.append({
            'goal': goal,
            'saved': saved_amount,
            'percent': percent,
            'remaining': max(0, goal.target_amount - saved_amount),
            'savings_breakdown': savings_breakdown,
            'status_color': 'danger' if percent < 40 else ('warning' if percent < 80 else 'success')
        })

    return render(request, 'finance/goal_list.html', {
        'goals': goals_data,
        'form': form
    })

@login_required
def goal_add_money(request, pk):
    goal = get_object_or_404(Goal, pk=pk, user=request.user)
    if request.method == 'POST':
        if goal.saved_amount >= goal.target_amount:
            messages.info(request, f'Goal "{goal.name}" is already completed!')
            return redirect('goal_list')
            
        try:
            amount = Decimal(request.POST.get('amount', 0))
            if amount > 0:
                goal.saved_amount += amount
                goal.save()
                messages.success(request, f'Succesfully added {request.user.currency_symbol}{amount} to {goal.name}!')
            else:
                messages.error(request, 'Please enter a valid amount.')
        except:
            messages.error(request, 'Invalid amount entered.')
    return redirect('goal_list')

@login_required
def goal_delete(request, pk):
    goal = get_object_or_404(Goal, pk=pk, user=request.user)
    if request.method == 'POST':
        goal.delete()
        messages.success(request, 'Goal deleted!')
        return redirect('goal_list')
    return redirect('goal_list')

@login_required
def challenge_list(request):
    active_challenge = Challenge.objects.filter(user=request.user, is_active=True).first()
    today = timezone.now().date()
    
    if active_challenge:
        # Check if challenge ended successfully
        if today > active_challenge.end_date:
            active_challenge.is_active = False
            active_challenge.is_successful = True
            active_challenge.save()
            
            # Award Badge
            badge_name = f"{active_challenge.challenge_type} Master"
            if active_challenge.challenge_type == 'Weekend': badge_name = "Weekend Warrior"
            elif active_challenge.challenge_type == 'Month': badge_name = "Month Legend"
            
            icon = 'fa-medal'
            if active_challenge.challenge_type == 'Day': icon = 'fa-calendar-check'
            elif active_challenge.challenge_type == 'Weekend': icon = 'fa-shield-halved'
            elif active_challenge.challenge_type == 'Month': icon = 'fa-crown'

            Badge.objects.get_or_create(
                user=request.user,
                name=badge_name,
                defaults={'icon': icon, 'description': f'Completed a {active_challenge.challenge_type} no-spend challenge!'}
            )
            messages.success(request, f'Congratulations! You completed the {active_challenge.challenge_type} challenge and earned a badge!')
            active_challenge = None # Clear for UI
        else:
            # Check for violations that might have happened since last visit (Manual imports etc)
            # Only count transactions created AFTER the challenge was started
            violations = Transaction.objects.filter(
                user=request.user,
                category__type='expense',
                date__range=[active_challenge.start_date, today],
                created_at__gte=active_challenge.created_at
            ).exists()
            
            if violations:
                active_challenge.is_active = False
                active_challenge.is_successful = False
                active_challenge.save()
                messages.error(request, f'Your {active_challenge.challenge_type} challenge was failed due to a spend.')
                active_challenge = None

    # Calculate Progress for active challenge
    progress_percent = 0
    days_left = 0
    if active_challenge:
        total_days = (active_challenge.end_date - active_challenge.start_date).days + 1
        elapsed_days = (today - active_challenge.start_date).days
        progress_percent = min(100, round((elapsed_days / total_days) * 100))
        days_left = max(0, (active_challenge.end_date - today).days)

    past_challenges = Challenge.objects.filter(user=request.user, is_active=False).order_by('-created_at')
    
    # Success Rate
    success_count = past_challenges.filter(is_successful=True).count()
    total_past = past_challenges.count()
    success_rate = round((success_count / total_past) * 100) if total_past > 0 else 0

    badges = Badge.objects.filter(user=request.user).order_by('-awarded_at')

    return render(request, 'finance/challenge_list.html', {
        'active_challenge': active_challenge,
        'past_challenges': past_challenges,
        'badges': badges,
        'progress_percent': progress_percent,
        'days_left': days_left,
        'success_rate': success_rate
    })

@login_required
def start_challenge(request):
    if request.method == 'POST':
        ctype = request.POST.get('type')
        start_date = timezone.now().date()
        
        if ctype == 'Day':
            end_date = start_date
        elif ctype == 'Weekend':
            # Weekend = until next Sunday
            days_to_sunday = (6 - start_date.weekday()) % 7
            if days_to_sunday == 0: days_to_sunday = 7 # If started on Sunday, next Sunday
            end_date = start_date + timedelta(days=days_to_sunday)
        elif ctype == 'Month':
            end_date = start_date + timedelta(days=30)
        else:
            return redirect('challenge_list')

        # Deactivate any previous active challenges
        Challenge.objects.filter(user=request.user, is_active=True).update(is_active=False, is_successful=False)
        
        Challenge.objects.create(
            user=request.user,
            challenge_type=ctype,
            start_date=start_date,
            end_date=end_date
        )
        messages.success(request, f'No-Spend {ctype} challenge started! Stay disciplined.')
        return redirect('challenge_list')
    
    return redirect('challenge_list')
